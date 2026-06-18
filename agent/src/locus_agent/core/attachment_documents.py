"""Office 附件（Excel / Word）文本抽取。"""

from __future__ import annotations

import io
from pathlib import PurePath

MAX_EXTRACT_CHARS = 16_000
MAX_SHEET_ROWS = 400
MAX_SHEET_COLS = 40

_OFFICE_EXTENSIONS = frozenset({".xlsx", ".xlsm", ".docx"})
_LEGACY_OFFICE_EXTENSIONS = frozenset({".xls", ".doc"})

_XLSX_MIMES = frozenset(
    {
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel.sheet.macroenabled.12",
    }
)
_DOCX_MIMES = frozenset(
    {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
)

OFFICE_EXTRACTED_MIME = "text/plain;charset=utf-8"
_OFFICE_FILENAME_EXTENSIONS = frozenset({".xlsx", ".xlsm", ".docx"})


def is_office_attachment(*, name: str, mime_type: str | None) -> bool:
    ext = PurePath(name or "").suffix.lower()
    if ext in _OFFICE_EXTENSIONS:
        return True
    mime = (mime_type or "").split(";", 1)[0].strip().lower()
    return mime in _XLSX_MIMES or mime in _DOCX_MIMES


def is_office_extracted_blob(*, name: str, mime_type: str | None) -> bool:
    """Office 解析成功后以 UTF-8 文本 blob 存储（避免重复解析）。"""
    mime = (mime_type or "").split(";", 1)[0].strip().lower()
    if mime != "text/plain":
        return False
    return PurePath(name or "").suffix.lower() in _OFFICE_FILENAME_EXTENSIONS


def is_office_binary_payload(data: bytes) -> bool:
    return len(data) >= 4 and data[:2] == b"PK"


def has_office_filename(name: str) -> bool:
    return PurePath(name or "").suffix.lower() in _OFFICE_FILENAME_EXTENSIONS


def is_legacy_office_attachment(*, name: str) -> bool:
    return PurePath(name or "").suffix.lower() in _LEGACY_OFFICE_EXTENSIONS


def _truncate(text: str) -> tuple[str, bool]:
    if len(text) <= MAX_EXTRACT_CHARS:
        return text, False
    return f"{text[:MAX_EXTRACT_CHARS]}\n...（文档过长，已截断）", True


def _extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows: list[str] = []
        for row_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
            if row_idx > MAX_SHEET_ROWS:
                rows.append(f"...（工作表 {sheet_name} 超过 {MAX_SHEET_ROWS} 行，已截断）")
                break
            cells = ["" if v is None else str(v).strip() for v in row[:MAX_SHEET_COLS]]
            if any(cells):
                rows.append("\t".join(cells))
        if rows:
            parts.append(f"## Sheet: {sheet_name}\n" + "\n".join(rows))
    wb.close()
    return "\n\n".join(parts).strip()


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n".join(parts).strip()


def extract_document_text(data: bytes, *, name: str, mime_type: str | None = None) -> tuple[str, bool, str | None]:
    """返回 (text, truncated, error)。"""
    if not data:
        return "", False, "empty file"
    if is_legacy_office_attachment(name=name):
        return "", False, "仅支持 .xlsx/.xlsm 与 .docx，请将 .xls/.doc 另存为新格式后上传"
    ext = PurePath(name or "").suffix.lower()
    mime = (mime_type or "").split(";", 1)[0].strip().lower()
    try:
        if ext in {".xlsx", ".xlsm"} or mime in _XLSX_MIMES:
            text = _extract_xlsx(data)
        elif ext == ".docx" or mime in _DOCX_MIMES:
            text = _extract_docx(data)
        else:
            return "", False, "不支持的 Office 格式"
    except Exception as exc:
        return "", False, f"文档解析失败: {exc}"
    if not text.strip():
        return "", False, "文档无可用文本内容"
    truncated_text, truncated = _truncate(text)
    return truncated_text, truncated, None
